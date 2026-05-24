"""
Module: 26_导出系统/DOCX导出/DOCX导出.py
Layer: 导出系统
Dependency: python-docx (optional), standard library
Called by: Export orchestrator in 26_导出系统
Purpose: Converts novel internal data structure into a .docx file with configurable styling.
         Implements pluggable exporter interface.
"""

import logging
import os
import json
from datetime import datetime
from typing import Any, Dict, List, Optional

# 尝试导入docx库，如果不可用则标记
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DOCXExporter:
    """可插拔的DOCX导出器，符合导出系统接口规范。"""

    # 默认配置
    DEFAULT_CONFIG = {
        "font_name": "Times New Roman",
        "font_size": 12,
        "title_font_size": 24,
        "heading_font_size": 16,
        "line_spacing": 1.5,
        "margin_top": Inches(1),
        "margin_bottom": Inches(1),
        "margin_left": Inches(1),
        "margin_right": Inches(1),
        "show_page_numbers": True,
        "output_dir": "output",
        "template_path": None,  # 可选的模板文件路径
        "style_options": {}  # 可扩展的样式选项
    }

    def __init__(self, config: Optional[Dict[str, Any]] = None, logger: Optional[logging.Logger] = None):
        """
        初始化DOCX导出器。
        
        Args:
            config: 自定义配置字典，将与默认配置合并。
            logger: 外部传入的logger，若未提供则使用模块级logger。
        """
        self.logger = logger or logging.getLogger(self.__class__.__name__)
        self.config = {**self.DEFAULT_CONFIG, **(config or {})}
        self._validate_config()
        self.logger.info("DOCXExporter initialized with config: %s", self.config)

    def _validate_config(self) -> None:
        """验证配置参数的有效性。"""
        if not isinstance(self.config.get("font_size"), (int, float)):
            raise ValueError("font_size must be a number")
        # 可扩展更多验证

    def export(self, novel_data: Dict[str, Any], output_filename: Optional[str] = None) -> str:
        """
        将小说数据导出为DOCX文件。
        
        Args:
            novel_data: 符合内部小说数据协议的结构化数据。
                必须包含字段: "title", "author", "chapters" (list of chapter dicts)
                每个chapter包含: "number", "title", "content" (字符串或段落列表)
                可选: "metadata", "settings"等
            output_filename: 输出文件名（不含扩展名）。若为None，则自动生成。
        
        Returns:
            生成的docx文件的绝对路径。
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is not installed. Please install it to use DOCX export.")

        self.logger.info("Starting DOCX export.")
        # 1. 验证输入数据
        self._validate_novel_data(novel_data)
        
        # 2. 创建文档
        doc = self._create_document()
        
        # 3. 构建内容
        self._build_title_page(doc, novel_data)
        self._build_chapters(doc, novel_data["chapters"])
        self._build_footer(doc, novel_data)
        
        # 4. 确定输出路径
        output_dir = self.config.get("output_dir", "output")
        os.makedirs(output_dir, exist_ok=True)
        if not output_filename:
            title_slug = novel_data.get("title", "novel").replace(" ", "_").lower()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{title_slug}_{timestamp}"
        filepath = os.path.join(output_dir, f"{output_filename}.docx")
        
        # 5. 保存
        doc.save(filepath)
        self.logger.info("DOCX exported successfully to: %s", filepath)
        return filepath

    def _validate_novel_data(self, data: Dict[str, Any]) -> None:
        """验证小说数据完整性，缺失必要字段时抛出异常。"""
        required = ["title", "author", "chapters"]
        for field in required:
            if field not in data:
                raise ValueError(f"Missing required field in novel data: {field}")
        for idx, ch in enumerate(data["chapters"]):
            if not isinstance(ch, dict) or "content" not in ch:
                raise ValueError(f"Chapter {idx} missing 'content' key")

    def _create_document(self) -> "Document":
        """创建并返回一个配置好的python-docx Document对象。"""
        if self.config.get("template_path") and os.path.exists(self.config["template_path"]):
            doc = Document(self.config["template_path"])
            self.logger.info("Using template: %s", self.config["template_path"])
        else:
            doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = self.config.get("font_name", "Times New Roman")
        font.size = Pt(self.config.get("font_size", 12))
        
        # 设置页边距
        for section in doc.sections:
            section.top_margin = self.config.get("margin_top", Inches(1))
            section.bottom_margin = self.config.get("margin_bottom", Inches(1))
            section.left_margin = self.config.get("margin_left", Inches(1))
            section.right_margin = self.config.get("margin_right", Inches(1))
        
        return doc

    def _build_title_page(self, doc: "Document", data: Dict[str, Any]) -> None:
        """构建标题页，包含书名、作者等信息。"""
        # 添加标题
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_paragraph.add_run(data.get("title", "Untitled"))
        run.font.size = Pt(self.config.get("title_font_size", 24))
        run.bold = True
        
        # 作者
        author_paragraph = doc.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_paragraph.add_run(f"By {data.get('author', 'Unknown')}")
        run.font.size = Pt(self.config.get("font_size", 12))
        
        # 可添加副标题、日期等（根据配置扩展）
        doc.add_page_break()

    def _build_chapters(self, doc: "Document", chapters: List[Dict[str, Any]]) -> None:
        """遍历所有章节并添加到文档中。"""
        for i, chapter in enumerate(chapters):
            self.logger.debug("Adding chapter %s", i+1)
            # 章节标题
            heading = doc.add_heading(level=1)
            run = heading.add_run(chapter.get("title", f"Chapter {i+1}"))
            run.font.size = Pt(self.config.get("heading_font_size", 16))
            
            # 章节内容
            content = chapter.get("content", "")
            if isinstance(content, str):
                paragraphs = content.split("\n")
            else:
                paragraphs = content  # 假设是列表
            
            for para_text in paragraphs:
                p = doc.add_paragraph(str(para_text))
                # 应用行距
                p.paragraph_format.line_spacing = self.config.get("line_spacing", 1.5)
            
            # 除了最后一章，添加分页符
            if i < len(chapters) - 1:
                doc.add_page_break()

    def _build_footer(self, doc: "Document", data: Dict[str, Any]) -> None:
        """构建页脚（如页码）根据配置显示。"""
        if self.config.get("show_page_numbers", True):
            for section in doc.sections:
                footer = section.footer
                if footer.is_linked_to_previous:
                    footer.is_linked_to_previous = False
                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 添加页码域代码
                run = p.add_run()
                fldChar1 = run._r.makeelement(0x3C, 0x3E)  # 简化，实际添加页码可用python-docx方式
                # 由于页码添加较复杂，这里用简单文本代替
                p.text = "Page "
    
    def shutdown(self) -> None:
        """优雅关闭，清理资源（目前无需特别操作）。"""
        self.logger.info("DOCXExporter shutdown.")
        pass

# ------------------ 自测试代码 ------------------
if __name__ == "__main__":
    # 配置日志输出到控制台
    logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    logger = logging.getLogger("DOCXExportTest")
    
    logger.info("Starting self-test for DOCX exporter.")
    
    if not DOCX_AVAILABLE:
        logger.error("python-docx not installed. Cannot run self-test.")
        raise SystemExit(1)
    
    # 模拟小说数据
    sample_novel = {
        "title": "The Adventures of Codex",
        "author": "AI Storyteller",
        "chapters": [
            {
                "number": 1,
                "title": "The Beginning",
                "content": "In a world full of code, there lived a curious variable named Codex.\nCodex longed to explore beyond its scope."
            },
            {
                "number": 2,
                "title": "The Loop",
                "content": "Codex entered a loop and met many iterations of itself.\nIt was a recursive nightmare."
            }
        ],
        "metadata": {"genre": "technology"}
    }
    
    # 自定义配置
    custom_config = {
        "font_name": "Arial",
        "font_size": 11,
        "output_dir": "test_output"
    }
    
    exporter = DOCXExporter(config=custom_config, logger=logger)
    try:
        output_path = exporter.export(sample_novel, output_filename="test_novel")
        logger.info(f"Self-test passed. File saved at: {output_path}")
    except Exception as e:
        logger.exception("Self-test failed with error: %s", e)
    finally:
        exporter.shutdown()